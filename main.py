import os
import cv2
import numpy as np
import pytesseract
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import layoutparser as lp
from PIL import Image

# Update the Tesseract OCR path if necessary
# For Windows users, it might be something like 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'
lang = 'heb'  # Hebrew language code

# Initialize the layout model for table detection
# We'll use a pre-trained model from Detectron2 via LayoutParser
model = lp.Detectron2LayoutModel(
    config_path='lp://PubLayNet/faster_rcnn_R_50_FPN_3x/config',
    label_map={0: "Text", 1: "Title", 2: "List", 3: "Table", 4: "Figure"},
    extra_config=["MODEL.ROI_HEADS.SCORE_THRESH_TEST", 0.5],
    device='cpu'  # Use 'cuda' if GPU is available
)

def image_contains_table(image_path):
    image = cv2.imread(image_path)
    image = image[..., ::-1]  # Convert BGR to RGB
    layout = model.detect(image)
    # Check if any detected blocks are tables
    table_blocks = [b for b in layout if b.type == 'Table']
    return len(table_blocks) > 0

def extract_tables_from_image(image_path):
    image = cv2.imread(image_path)
    image_rgb = image[..., ::-1]  # Convert BGR to RGB
    layout = model.detect(image_rgb)
    # Filter out table blocks
    table_blocks = [b for b in layout if b.type == 'Table']
    tables_data = []
    for idx, table_block in enumerate(table_blocks):
        # Crop the table region
        x_1, y_1, x_2, y_2 = map(int, table_block.coordinates)
        table_image = image[y_1:y_2, x_1:x_2]
        # Process the table image to extract text
        table_data = extract_table_data(table_image)
        tables_data.append((table_block, table_data))
    return tables_data

def extract_table_data(table_image):
    # Preprocess the table image if necessary
    # Convert to grayscale
    gray = cv2.cvtColor(table_image, cv2.COLOR_BGR2GRAY)
    # Optional: Apply thresholding or other preprocessing steps
    # Perform OCR on the table image
    ocr_result = pytesseract.image_to_data(
        gray,
        lang=lang,
        config='--psm 6',  # Assume a single uniform block of text
        output_type=pytesseract.Output.DICT
    )
    # Organize OCR results into a table structure
    n_boxes = len(ocr_result['level'])
    data = []
    for i in range(n_boxes):
        if int(ocr_result['conf'][i]) > 60:
            (x, y, w, h, text) = (
                ocr_result['left'][i],
                ocr_result['top'][i],
                ocr_result['width'][i],
                ocr_result['height'][i],
                ocr_result['text'][i]
            )
            data.append({
                'text': text,
                'x': x,
                'y': y,
                'w': w,
                'h': h
            })
    # Group text by rows based on y-coordinate
    rows = {}
    for d in data:
        y_center = d['y'] + d['h'] / 2
        assigned = False
        for row_y in rows:
            if abs(y_center - row_y) < 10:
                rows[row_y].append(d)
                assigned = True
                break
        if not assigned:
            rows[y_center] = [d]
    # Sort rows by y-coordinate (top to bottom)
    sorted_rows = sorted(rows.items(), key=lambda x: x[0])
    table_data = []
    for row_y, texts in sorted_rows:
        # Sort texts in the row by x-coordinate (right to left for Hebrew)
        sorted_texts = sorted(texts, key=lambda x: -x['x'])
        row_data = [t['text'] for t in sorted_texts]
        table_data.append(row_data)
    return table_data

def replace_image_with_tables(paragraph, run, tables_data):
    # Remove the run containing the image
    run._element.getparent().remove(run._element)
    # Insert each table in place of the image
    for table_block, table_data in tables_data:
        max_cols = max(len(row) for row in table_data)
        doc_table = paragraph.add_table(rows=0, cols=max_cols)
        doc_table.style = 'Table Grid'
        for row_data in table_data:
            # Ensure the row has the correct number of columns
            row_cells = doc_table.add_row().cells
            for idx in range(max_cols):
                cell = row_cells[idx]
                if idx < len(row_data):
                    cell_text = row_data[idx]
                else:
                    cell_text = ''
                cell.text = cell_text
                # Apply right-to-left text direction for Hebrew
                cell.paragraphs[0].paragraph_format.right_to_left = True
                # Set font size
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(12)
        # Add a paragraph break between tables if there are multiple tables
        paragraph = paragraph.insert_paragraph_before()

def process_document(input_file, output_file):
    document = Document(input_file)
    for paragraph in document.paragraphs:
        runs_to_remove = []
        for run in paragraph.runs:
            if 'graphic' in run._element.xml:
                drawing_elements = run._element.xpath('.//w:drawing')
                for drawing in drawing_elements:
                    blip = drawing.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                    if blip:
                        embed = blip[0].get(qn('r:embed'))
                        image_part = document.part.related_parts[embed]
                        with open('temp_image.png', 'wb') as f:
                            f.write(image_part.blob)
                        image_filename = 'temp_image.png'
                        # Check if the image contains any tables
                        if image_contains_table(image_filename):
                            # Extract tables from the image
                            tables_data = extract_tables_from_image(image_filename)
                            if tables_data:
                                replace_image_with_tables(paragraph, run, tables_data)
                            else:
                                raise ValueError("Failed to extract table data accurately.")
                        else:
                            # Keep the image as is
                            pass
                        os.remove(image_filename)
        # Remove processed runs
        for run in runs_to_remove:
            run._element.getparent().remove(run._element)
    document.save(output_file)

if __name__ == "__main__":
    input_docx = 'input.docx'    # Replace with your input file path
    output_docx = 'output.docx'  # Replace with your desired output file path
    process_document(input_docx, output_docx)
